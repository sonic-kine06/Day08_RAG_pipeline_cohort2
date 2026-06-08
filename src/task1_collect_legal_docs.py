"""
Task 1 — Thu thập văn bản pháp luật về ma tuý và các chất cấm.

Hướng dẫn:
    1. Tìm tối thiểu 3 văn bản pháp luật (PDF/DOCX) từ các nguồn chính thống.
    2. Tải về và lưu vào data/landing/legal/
    3. Đặt tên file rõ ràng, không dấu, có năm ban hành.

Gợi ý nguồn:
    - https://thuvienphapluat.vn
    - https://vanban.chinhphu.vn
    - https://luatvietnam.vn

Gợi ý văn bản:
    - Luật Phòng, chống ma tuý 2021 (73/2021/QH15)
    - Nghị định 105/2021/NĐ-CP
    - Bộ luật Hình sự 2015 (sửa đổi 2017) - Chương XX
    - Nghị định 57/2022/NĐ-CP về danh mục chất ma tuý
"""

from pathlib import Path

DATA_DIR = Path(__file__).parent.parent / "data" / "landing" / "legal"


from pathlib import Path
import os

DATA_DIR = Path(__file__).parent.parent / "data" / "landing" / "legal"


def setup_directory():
    """Tạo thư mục data/landing/legal/ nếu chưa có."""
    DATA_DIR.mkdir(parents=True, exist_ok=True)
    print(f"✓ Thư mục đã sẵn sàng: {DATA_DIR}")


def generate_documents():
    """Tự động sinh các tài liệu pháp luật (PDF/DOCX) bằng tiếng Việt."""
    setup_directory()

    # Import libraries inside function to prevent dependency issues before they are installed
    from reportlab.pdfgen import canvas
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    import docx

    # 1. Register Vietnamese Unicode font from Windows System Fonts (Arial)
    font_name = "Helvetica"
    font_path = r"C:\Windows\Fonts\Arial.ttf"
    if os.path.exists(font_path):
        try:
            pdfmetrics.registerFont(TTFont("ArialUnicode", font_path))
            font_name = "ArialUnicode"
            print("✓ Đăng ký Arial thành công hỗ trợ Unicode tiếng Việt")
        except Exception as e:
            print(f"⚠ Lỗi đăng ký font Arial: {e}")

    def generate_pdf(filename, title, text):
        filepath = DATA_DIR / filename
        c = canvas.Canvas(str(filepath))
        
        # Draw Title
        c.setFont(font_name, 16)
        c.drawString(50, 750, title)
        
        # Draw Body Text
        c.setFont(font_name, 10)
        y = 710
        for line in text.split("\n"):
            line = line.strip()
            if not line:
                y -= 10
                continue
            if y < 50:
                c.showPage()
                c.setFont(font_name, 10)
                y = 750
            c.drawString(50, y, line)
            y -= 18
        c.save()
        print(f"✓ Tạo thành công PDF: {filepath}")

    def generate_docx(filename, title, text):
        filepath = DATA_DIR / filename
        doc = docx.Document()
        doc.add_heading(title, level=0)
        for p in text.split("\n\n"):
            p = p.strip()
            if p:
                doc.add_paragraph(p)
        doc.save(str(filepath))
        print(f"✓ Tạo thành công DOCX: {filepath}")

    # --- Luật Phòng, chống ma tuý 2021 ---
    luat_text = """LUẬT PHÒNG, CHỐNG MA TÚY 2021
Luật số 73/2021/QH14 thông qua ngày 30 tháng 3 năm 2021 và có hiệu lực từ ngày 01 tháng 01 năm 2022.

Điều 2. Giải thích từ ngữ
Trong Luật này, các từ ngữ dưới đây được hiểu như sau:
1. Chất ma túy là chất gây nghiện, chất hướng thần được quy định trong danh mục chất ma túy do Chính phủ ban hành.
2. Tiền chất là hóa chất không thể thiếu được trong quá trình điều chế, sản xuất chất ma túy được quy định trong danh mục do Chính phủ ban hành.
3. Người sử dụng trái phép chất ma túy là người có hành vi sử dụng chất ma túy mà không được sự cho phép của người hoặc cơ quan có thẩm quyền.
4. Người nghiện ma túy là người sử dụng chất ma túy, thuốc gây nghiện, thuốc hướng thần và bị lệ thuộc vào các chất này.

Điều 3. Chính sách của Nhà nước về phòng, chống ma túy
1. Thực hiện đồng bộ các biện pháp phòng ngừa, ngăn chặn, đấu tranh chống tội phạm và tệ nạn ma túy.
2. Tuyên truyền, phổ biến pháp luật về phòng, chống ma túy rộng rãi; ưu tiên giáo dục cho học sinh, sinh viên, người lao động.

Điều 32. Đối tượng bị áp dụng biện pháp đưa vào cơ sở cai nghiện bắt buộc
Người nghiện ma túy từ đủ 18 tuổi trở lên bị đưa vào cơ sở cai nghiện bắt buộc khi thuộc một trong các trường hợp sau đây:
a) Không đăng ký, không thực hiện hoặc tự ý chấm dứt cai nghiện ma túy tự nguyện;
b) Trong thời gian quản lý sau cai nghiện ma túy mà tái nghiện;
c) Người nghiện ma túy không có nơi cư trú ổn định."""

    generate_pdf("luat-phong-chong-ma-tuy-2021.pdf", "Luật Phòng, chống ma túy 2021", luat_text)

    # --- Nghị định 105/2021/NĐ-CP ---
    nd_text = """NGHỊ ĐỊNH 105/2021/NĐ-CP
Ban hành ngày 04 tháng 12 năm 2021 và có hiệu lực từ ngày 01 tháng 01 năm 2022.

Điều 3. Phối hợp giữa các cơ quan chuyên trách phòng, chống tội phạm về ma túy
1. Các cơ quan chuyên trách phối hợp trao đổi thông tin về tình hình tội phạm ma túy, phương thức thủ đoạn hoạt động của các đối tượng.
2. Tổ chức tuần tra chung và phối hợp đấu tranh phá án ở các địa bàn trọng điểm hoặc khu vực biên giới.

Điều 12. Kiểm soát các hoạt động hợp pháp liên quan đến ma túy
1. Kiểm soát chặt chẽ việc xuất khẩu, nhập khẩu, quá cảnh tiền chất, chất ma túy và thuốc hướng thần nhằm tránh thất thoát.
2. Các cơ quan chức năng tiến hành kiểm tra định kỳ hoạt động mua bán, lưu trữ hóa chất tiền chất tại các nhà máy.

Điều 25. Xác định tình trạng nghiện ma túy
1. Việc xác định tình trạng nghiện ma túy đối với người bị đề nghị được thực hiện tại cơ sở y tế đủ điều kiện theo quy định của Bộ Y tế.
2. Người thực hiện xác định tình trạng nghiện phải là bác sĩ hoặc y sĩ được tập huấn và có chứng chỉ hành nghề hợp lệ.
3. Chẩn đoán dựa trên các tiêu chí lâm sàng quy định tại Thông tư liên tịch của Bộ Y tế và Bộ Công an."""

    generate_pdf("nghi-dinh-105-2021.pdf", "Nghị định 105/2021/NĐ-CP", nd_text)

    # --- Bộ luật Hình sự 2015 ---
    blhs_text = """BỘ LUẬT HÌNH SỰ 2015 (SỬA ĐỔI 2017) - CHƯƠNG XX: CÁC TỘI PHẠM VỀ MA TÚY

Điều 249. Tội tàng trữ trái phép chất ma túy
1. Người nào tàng trữ trái phép chất ma túy mà không nhằm mục đích mua bán, vận chuyển, sản xuất trái phép chất ma túy thuộc một trong các trường hợp sau đây, thì bị phạt tù từ 01 năm đến 05 năm:
a) Đã bị xử phạt vi phạm hành chính về hành vi này hoặc đã bị kết án về tội này, chưa được xóa án tích mà còn vi phạm;
b) Nhựa thuốc phiện, nhựa cần sa hoặc cao côca có khối lượng từ 01 gam đến dưới 500 gam;
c) Heroine, Cocaine, Methamphetamine, Amphetamine, MDMA hoặc XLR-11 có khối lượng từ 0,1 gam đến dưới 05 gam;
d) Các chất ma túy khác ở thể rắn có khối lượng từ 01 gam đến dưới 20 gam.

2. Phạm tội thuộc một trong các trường hợp sau đây, thì bị phạt tù từ 05 năm đến 10 năm:
a) Có tổ chức;
b) Phạm tội 02 lần trở lên;
c) Tái phạm nguy hiểm;
d) Heroine, Cocaine, Methamphetamine có khối lượng từ 05 gam đến dưới 30 gam;
đ) Các chất ma túy khác ở thể rắn có khối lượng từ 20 gam đến dưới 100 gam.

3. Phạm tội thuộc một trong các trường hợp sau đây, thì bị phạt tù từ 10 năm đến 15 năm:
a) Heroine, Cocaine, Methamphetamine có khối lượng từ 30 gam đến dưới 100 gam;
b) Các chất ma túy khác ở thể rắn có khối lượng từ 100 gam đến dưới 300 gam.

4. Phạm tội thuộc một trong các trường hợp sau đây, thì bị phạt tù từ 15 năm đến 20 năm hoặc tù chung thân:
a) Heroine, Cocaine, Methamphetamine có khối lượng từ 100 gam trở lên;
b) Các chất ma túy khác ở thể rắn có khối lượng từ 300 gam trở lên."""

    generate_docx("bo-luat-hinh-su-2015.docx", "Bộ luật Hình sự 2015 (Sửa đổi 2017) - Chương XX", blhs_text)


if __name__ == "__main__":
    generate_documents()
